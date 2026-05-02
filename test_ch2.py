from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, shutil

doc = Document('EDR-Documentaion-Chapter-1-V1.docx')

# Clear all paragraphs
for p in doc.paragraphs:
    p.clear()
for t in doc.tables:
    t._element.getparent().remove(t._element)

print("base loaded")
